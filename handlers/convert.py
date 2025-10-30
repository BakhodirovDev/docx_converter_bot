from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph


def analyze_docx_file(docx_path: str) -> dict:
    """
    DOCX faylni tahlil qilib, rasm, formula va boshqa muammolarni aniqlaydi.
    
    Returns:
        {
            'total_questions': int,
            'has_images': bool,
            'has_equations': bool,
            'image_count': int,
            'equation_count': int,
            'problematic_questions': list  # Rasm/formula bor savol raqamlari
        }
    """
    doc = Document(docx_path)
    
    total_questions = 0
    has_images = False
    has_equations = False
    image_count = 0
    equation_count = 0
    problematic_questions = []
    
    question_num = 0
    
    for table in doc.tables:
        for row in table.rows:
            question = row.cells[0].text.strip()
            
            if question:
                question_num += 1
                total_questions += 1
                
                # Har bir katakda rasm yoki formula tekshirish
                question_has_problem = False
                
                for cell in row.cells:
                    # Rasmlarni tekshirish
                    for paragraph in cell.paragraphs:
                        # Inline shapes (rasmlar)
                        if paragraph._element.xpath('.//w:drawing'):
                            has_images = True
                            image_count += 1
                            question_has_problem = True
                        
                        # OLE objects (formula kabi)
                        if paragraph._element.xpath('.//w:object'):
                            has_equations = True
                            equation_count += 1
                            question_has_problem = True
                        
                        # Math equations (MathML)
                        if paragraph._element.xpath('.//m:oMath'):
                            has_equations = True
                            equation_count += 1
                            question_has_problem = True
                
                if question_has_problem:
                    problematic_questions.append(question_num)
    
    return {
        'total_questions': total_questions,
        'has_images': has_images,
        'has_equations': has_equations,
        'image_count': image_count,
        'equation_count': equation_count,
        'problematic_questions': problematic_questions
    }


def convert_docx_to_txt(docx_path: str, txt_path: str) -> str:
    """
    DOCX fayldan savol-javoblarni extract qilib TXT ga yozadi.
    
    Format:
    ? Savol matni
    + To'g'ri javob
    - Noto'g'ri javob 1
    - Noto'g'ri javob 2
    """
    doc = Document(docx_path)
    
    with open(txt_path, 'w', encoding='utf-8') as file:
        # Jadvallardagi savol-javoblarni extract qilish
        for table in doc.tables:
            for row in table.rows:
                # Birinchi ustun - savol
                question = row.cells[0].text.strip()
                
                # Qolgan ustunlar - javoblar
                answers = [cell.text.strip() for cell in row.cells[1:]]
                
                if question:  # Faqat savol bo'lsa yozamiz
                    file.write(f"? {question}\n")
                    
                    # Birinchi javob - to'g'ri javob (+)
                    if len(answers) > 0 and answers[0]:
                        file.write(f"+ {answers[0]}\n")
                    
                    # Qolgan javoblar - noto'g'ri javoblar (-)
                    for ans in answers[1:]:
                        if ans:  # Bo'sh bo'lmasa
                            file.write(f"- {ans}\n")
                    
                    file.write("\n")  # Har bir savol-javobdan keyin bo'sh qator
    
    return txt_path
